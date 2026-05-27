#!/usr/bin/env python3
# /// script
# requires-python = ">=3.11"
# dependencies = ["python-docx>=1.1.2"]
# ///
"""create_docx — build a formatted Microsoft Word (.docx) document from a JSON spec.

The agent writes a structured spec (title + sections of headings, paragraphs,
bullets, numbered lists, and tables); this renders a clean, styled .docx.

Spec schema (JSON):
{
  "title":    "Quarterly Report",        # optional — renders a title block
  "subtitle": "Q2 2026",                 # optional
  "author":   "Flopsy",                  # optional — small line under the title
  "accent":   "1F4E35",                  # optional hex (no #), used for headings + table header fill
  "sections": [
    {
      "heading":   "Summary",            # optional
      "level":     1,                     # 1-3, default 1
      "paragraphs": ["First para.", "Second."],
      "bullets":    ["point a", "point b"],
      "numbered":   ["step 1", "step 2"],
      "table": { "headers": ["Name","Value"], "rows": [["x","1"],["y","2"]] }
    }
  ]
}

Any section field is optional; include only what you need. Order within a section
is: heading -> paragraphs -> bullets -> numbered -> table.

Examples:
  uv run create_docx.py --spec /workspace/work/scratch/report.json --output /workspace/work/exports/report.docx
  echo '{"title":"Hi","sections":[{"paragraphs":["Body text."]}]}' | uv run create_docx.py --stdin --output /workspace/work/exports/hi.docx
  uv run create_docx.py --from-markdown /workspace/work/scratch/notes.md --output /workspace/work/exports/notes.docx   # requires pandoc

Output: single-line JSON on stdout -> {"output_path": "...", "format": "docx", "bytes": 12345, "sections": 3}
Exit codes: 0 ok | 2 usage/spec error | 3 render error
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path


def _die(msg: str, code: int = 2) -> "None":
    print(f"create_docx: {msg}", file=sys.stderr)
    raise SystemExit(code)


def _load_spec(args: argparse.Namespace) -> dict:
    raw: str
    if args.stdin:
        raw = sys.stdin.read()
    elif args.spec_json is not None:
        raw = args.spec_json
    elif args.spec is not None:
        try:
            raw = Path(args.spec).read_text(encoding="utf-8")
        except OSError as e:
            _die(f"cannot read --spec file: {e}")
    else:
        _die("provide one of --spec FILE, --spec-json STR, --stdin, or --from-markdown")
    try:
        spec = json.loads(raw)
    except json.JSONDecodeError as e:
        _die(f"spec is not valid JSON: {e}")
    if not isinstance(spec, dict):
        _die("spec must be a JSON object")
    return spec


def _hex_to_rgb(value: str):
    from docx.shared import RGBColor

    v = value.strip().lstrip("#")
    if len(v) != 6:
        _die(f"accent must be a 6-digit hex color, got {value!r}")
    try:
        return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))
    except ValueError:
        _die(f"accent is not valid hex: {value!r}")


def _shade_cell(cell, hex_fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _render(spec: dict, output: Path) -> dict:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    accent_hex = str(spec.get("accent", "1F4E35")).strip().lstrip("#")
    accent = _hex_to_rgb(accent_hex)

    doc = Document()

    title = spec.get("title")
    if title:
        p = doc.add_paragraph()
        run = p.add_run(str(title))
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = accent
    subtitle = spec.get("subtitle")
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(str(subtitle))
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    author = spec.get("author")
    if author:
        p = doc.add_paragraph()
        run = p.add_run(str(author))
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    if title or subtitle or author:
        doc.add_paragraph()

    sections = spec.get("sections", [])
    if not isinstance(sections, list):
        _die("`sections` must be an array")

    for idx, sec in enumerate(sections):
        if not isinstance(sec, dict):
            _die(f"section[{idx}] must be an object")

        heading = sec.get("heading")
        if heading:
            level = sec.get("level", 1)
            try:
                level = max(1, min(3, int(level)))
            except (TypeError, ValueError):
                level = 1
            h = doc.add_heading(str(heading), level=level)
            for run in h.runs:
                run.font.color.rgb = accent

        for para in sec.get("paragraphs", []) or []:
            doc.add_paragraph(str(para))

        for item in sec.get("bullets", []) or []:
            doc.add_paragraph(str(item), style="List Bullet")

        for item in sec.get("numbered", []) or []:
            doc.add_paragraph(str(item), style="List Number")

        table = sec.get("table")
        if table:
            headers = table.get("headers") or []
            rows = table.get("rows") or []
            ncols = len(headers) if headers else (len(rows[0]) if rows else 0)
            if ncols == 0:
                _die(f"section[{idx}].table has no columns")
            t = doc.add_table(rows=0, cols=ncols)
            t.style = "Table Grid"
            if headers:
                hdr = t.add_row().cells
                for c, text in enumerate(headers):
                    if c >= ncols:
                        break
                    _shade_cell(hdr[c], accent_hex)
                    cell_p = hdr[c].paragraphs[0]
                    run = cell_p.add_run(str(text))
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for row in rows:
                cells = t.add_row().cells
                for c in range(ncols):
                    cells[c].text = str(row[c]) if c < len(row) else ""

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return {"sections": len(sections)}


def _from_markdown(md_path: Path, output: Path) -> dict:
    if shutil.which("pandoc") is None:
        _die("--from-markdown needs pandoc, which is not installed", code=3)
    if not md_path.exists():
        _die(f"markdown file not found: {md_path}")
    output.parent.mkdir(parents=True, exist_ok=True)
    proc = subprocess.run(
        ["pandoc", str(md_path), "-o", str(output)],
        capture_output=True,
        text=True,
    )
    if proc.returncode != 0:
        _die(f"pandoc failed: {proc.stderr.strip()}", code=3)
    return {"sections": 0, "engine": "pandoc"}


def main(argv: "list[str] | None" = None) -> int:
    ap = argparse.ArgumentParser(
        prog="create_docx",
        description="Build a formatted .docx from a JSON spec (see module docstring for the schema).",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    src = ap.add_mutually_exclusive_group()
    src.add_argument("--spec", metavar="FILE", help="Path to a JSON spec file.")
    src.add_argument("--spec-json", metavar="STR", help="Inline JSON spec string.")
    src.add_argument("--stdin", action="store_true", help="Read the JSON spec from stdin.")
    src.add_argument("--from-markdown", metavar="FILE", help="Convert a Markdown file to .docx via pandoc (ignores the JSON spec).")
    ap.add_argument("-o", "--output", required=True, metavar="PATH", help="Output .docx path (parent dirs auto-created).")
    args = ap.parse_args(argv)

    output = Path(args.output)
    if output.suffix.lower() != ".docx":
        output = output.with_suffix(".docx")

    if args.from_markdown:
        extra = _from_markdown(Path(args.from_markdown), output)
    else:
        spec = _load_spec(args)
        extra = _render(spec, output)

    result = {
        "output_path": str(output.resolve()),
        "format": "docx",
        "bytes": output.stat().st_size if output.exists() else 0,
        **extra,
    }
    print(json.dumps(result))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
